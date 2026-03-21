"""
Extracts plain text from PDF, DOCX, and TXT uploads.
Uses pdfplumber for PDF (handles multi-column layouts),
python-docx for Word documents.
"""
import io
import logging
from pathlib import Path
from typing import Optional

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)

# Demo documents used when no file is uploaded
DEMO_RESUME = """
John Smith — Senior Software Developer
Skills: Python, JavaScript, React, SQL, Git, Linux, REST APIs, Docker (basic)
Experience:
  - 4 years backend development at TechCorp (Python, Django, PostgreSQL)
  - 2 years frontend at StartupXYZ (React, TypeScript, Node.js)
Education: B.Sc. Computer Science, 2018
Certifications: AWS Cloud Practitioner (expired 2022)
"""

DEMO_JD = {
    "Software Engineer": """
    Senior Software Engineer — CloudNative Inc.
    Required: Python, System Design, Kubernetes, Docker, CI/CD pipelines,
    distributed systems, microservices, AWS/GCP, Terraform, REST APIs,
    Git, SQL, testing frameworks (pytest, Jest), code review.
    Nice-to-have: Go, GraphQL, Kafka, Prometheus.
    """,
    "Data Scientist": """
    Data Scientist — MLOps Focus — DataDriven Co.
    Required: Python, Statistics, Machine Learning, Deep Learning (PyTorch/TensorFlow),
    MLOps, Kubeflow, Feature Engineering, A/B Testing, Spark, SQL,
    model deployment, data pipelines, Pandas, Scikit-learn.
    """,
    "Product Manager": """
    Senior Product Manager — GrowthTech Ltd.
    Required: Product Strategy, OKR Framework, Agile/Scrum, Roadmapping,
    Stakeholder Management, GTM Strategy, Data Analytics, SQL (basic),
    UX Research, Competitive Analysis, JIRA, Figma collaboration.
    """,
}


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Route to the correct extractor based on file extension."""
    suffix = Path(filename).suffix.lower()
    try:
        if suffix == ".pdf":
            return _extract_pdf(file_bytes)
        elif suffix in (".docx", ".doc"):
            return _extract_docx(file_bytes)
        elif suffix == ".txt":
            return file_bytes.decode("utf-8", errors="replace")
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    except Exception as e:
        logger.error("File extraction failed for %s: %s", filename, e)
        raise


def _extract_pdf(data: bytes) -> str:
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            # crop to remove headers/footers (top/bottom 5%)
            h = page.height
            cropped = page.crop((0, h * 0.05, page.width, h * 0.95))
            raw = cropped.extract_text(x_tolerance=2, y_tolerance=2)
            if raw:
                text_parts.append(raw)
    return "\n".join(text_parts)


def _extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # also grab table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def get_demo_resume() -> str:
    return DEMO_RESUME.strip()


def get_demo_jd(role: str) -> str:
    return DEMO_JD.get(role, DEMO_JD["Software Engineer"]).strip()


def validate_file_size(size_bytes: int, max_mb: int = 10) -> Optional[str]:
    """Returns an error string if file is too large, else None."""
    if size_bytes > max_mb * 1024 * 1024:
        return f"File too large: {size_bytes / 1024 / 1024:.1f} MB (max {max_mb} MB)"
    return None